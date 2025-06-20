from typing import Dict, List, Optional
import json
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

class ReportGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        
    def generate(self, results: Dict, report_type: str, 
                include_options: Dict, format: str) -> bytes:
        """Genera reporte en el formato especificado"""
        
        if format == "PDF":
            return self._generate_pdf(results, report_type, include_options)
        elif format == "HTML":
            return self._generate_html(results, report_type, include_options)
        elif format == "DOCX":
            return self._generate_docx(results, report_type, include_options)
        elif format == "JSON":
            return self._generate_json(results, report_type, include_options)
    
    def _generate_pdf(self, results: Dict, report_type: str, options: Dict) -> bytes:
        """Genera reporte PDF"""
        from io import BytesIO
        buffer = BytesIO()
        
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1E3A8A')
        )
        
        story.append(Paragraph("Reporte de Comparación de PDFs", title_style))
        story.append(Spacer(1, 12))
        
        # Fecha
        story.append(Paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 
                             self.styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Resumen ejecutivo
        if options.get('summary', True):
            story.append(Paragraph("Resumen Ejecutivo", self.styles['Heading2']))
            summary = self._generate_summary(results)
            story.append(Paragraph(summary, self.styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Resultados principales
        story.append(Paragraph("Resultados del Análisis", self.styles['Heading2']))
        
        # Tabla de métricas
        data = [['Métrica', 'Valor']]
        
        if 'basic' in results:
            data.append(['Similitud General', f"{results['basic']['similarity_ratio']:.1%}"])
            data.append(['Líneas Añadidas', str(results['basic']['added_lines'])])
            data.append(['Líneas Eliminadas', str(results['basic']['removed_lines'])])
        
        if 'semantic' in results:
            data.append(['Similitud Semántica', f"{results['semantic']['overall_similarity']:.1%}"])
        
        if 'tfidf' in results:
            data.append(['Similitud TF-IDF', f"{results['tfidf']['cosine_similarity']:.1%}"])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 12))
        
        # Recomendaciones
        if options.get('recommendations', True):
            story.append(Paragraph("Recomendaciones", self.styles['Heading2']))
            recommendations = self._generate_recommendations(results)
            for rec in recommendations:
                story.append(Paragraph(f"• {rec}", self.styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _generate_html(self, results: Dict, report_type: str, options: Dict) -> str:
        """Genera reporte HTML"""
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Reporte de Comparación PDF</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #1E3A8A; }}
                h2 {{ color: #374151; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #4B5563; color: white; }}
                .metric {{ background-color: #F3F4F6; padding: 10px; margin: 10px 0; }}
            </style>
        </head>
        <body>
            <h1>Reporte de Comparación de PDFs</h1>
            <p>Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        """
        
        if options.get('summary', True):
            html += f"""
            <h2>Resumen Ejecutivo</h2>
            <p>{self._generate_summary(results)}</p>
            """
        
        # Métricas
        html += "<h2>Resultados del Análisis</h2>"
        html += "<table>"
        html += "<tr><th>Métrica</th><th>Valor</th></tr>"
        
        if 'basic' in results:
            html += f"<tr><td>Similitud General</td><td>{results['basic']['similarity_ratio']:.1%}</td></tr>"
        
        if 'semantic' in results:
            html += f"<tr><td>Similitud Semántica</td><td>{results['semantic']['overall_similarity']:.1%}</td></tr>"
        
        html += "</table>"
        
        if options.get('recommendations', True):
            html += "<h2>Recomendaciones</h2><ul>"
            for rec in self._generate_recommendations(results):
                html += f"<li>{rec}</li>"
            html += "</ul>"
        
        html += "</body></html>"
        
        return html.encode('utf-8')
    
    def _generate_docx(self, results: Dict, report_type: str, options: Dict) -> bytes:
        """Genera reporte DOCX"""
        doc = Document()
        
        # Título
        doc.add_heading('Reporte de Comparación de PDFs', 0)
        doc.add_paragraph(f'Fecha: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        # Resumen
        if options.get('summary', True):
            doc.add_heading('Resumen Ejecutivo', level=1)
            doc.add_paragraph(self._generate_summary(results))
        
        # Resultados
        doc.add_heading('Resultados del Análisis', level=1)
        
        # Tabla de resultados
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        if 'basic' in results:
            row_cells = table.add_row().cells
            row_cells[0].text = 'Similitud General'
            row_cells[1].text = f"{results['basic']['similarity_ratio']:.1%}"
        
        # Guardar en buffer
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _generate_json(self, results: Dict, report_type: str, options: Dict) -> bytes:
        """Genera reporte JSON"""
        report = {
            'metadata': {
                'report_type': report_type,
                'generated_at': datetime.now().isoformat(),
                'options': options
            },
            'summary': self._generate_summary(results) if options.get('summary', True) else None,
            'results': results,
            'recommendations': self._generate_recommendations(results) if options.get('recommendations', True) else []
        }
        
        return json.dumps(report, indent=2, ensure_ascii=False).encode('utf-8')
    
    def _generate_summary(self, results: Dict) -> str:
        """Genera resumen ejecutivo"""
        similarity = results.get('basic', {}).get('similarity_ratio', 0)
        
        if similarity > 0.9:
            level = "muy alta"
        elif similarity > 0.7:
            level = "alta"
        elif similarity > 0.5:
            level = "moderada"
        else:
            level = "baja"
        
        return f"""
        Los documentos analizados muestran una similitud {level} ({similarity:.1%}).
        Se detectaron {results.get('basic', {}).get('added_lines', 0)} líneas añadidas y
        {results.get('basic', {}).get('removed_lines', 0)} líneas eliminadas.
        El análisis semántico revela una similitud conceptual del 
        {results.get('semantic', {}).get('overall_similarity', 0):.1%}.
        """
    
    def _generate_recommendations(self, results: Dict) -> List[str]:
        """Genera recomendaciones basadas en los resultados"""
        recommendations = []
        
        similarity = results.get('basic', {}).get('similarity_ratio', 0)
        
        if similarity < 0.5:
            recommendations.append("Los documentos son significativamente diferentes. Revisar cambios mayores.")
        elif similarity < 0.8:
            recommendations.append("Se detectaron cambios moderados. Verificar secciones modificadas.")
        else:
            recommendations.append("Los documentos son muy similares. Revisar cambios menores.")
        
        if results.get('basic', {}).get('added_lines', 0) > 50:
            recommendations.append("Se agregó contenido sustancial. Revisar nuevas secciones.")
        
        if results.get('semantic', {}).get('unique_chunks_doc2', []):
            recommendations.append("El segundo documento contiene conceptos nuevos no presentes en el primero.")
        
        return recommendations